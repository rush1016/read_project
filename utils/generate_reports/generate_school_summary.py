from django.http import HttpResponse

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

from students.models import Student
from read_app.models import School, ClassSection

from utils.srp_setup import SrpSetup
from utils.table_setup_utils import TableSetupUtils

def generate_school_reading_profile(school_id):
    # Classes per grade and section
    school = School.objects.get(pk=school_id)
    all_students = Student.objects.filter(school=school)
    all_sections = ClassSection.objects.all().order_by('section_name')

    # Setup
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(12)

    title = doc.add_paragraph('School Reading Profile')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(13)

    info_line1 = doc.add_paragraph('School: ')
    info_line1.add_run(school.name).underline = True
    info_line1.add_run("        Division: ")
    info_line1.add_run(school.division).underline = True
    info1_format = info_line1.paragraph_format
    info1_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    info_line2 = doc.add_paragraph('District: ')
    info_line2.add_run(school.district).underline = True
    info_line2.add_run("        Region: ")
    info_line2.add_run(school.region).underline = True
    info2_format = info_line2.paragraph_format
    info2_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    doc.add_paragraph()

    table = doc.add_table(rows=24, cols=5)
    table.style = 'Table Grid'
    hdr_row = table.rows[0]
    hdr_row2 = table.rows[1]

    hdr_row.cells[0].text = 'Grade'
    hdr_row.cells[1].text = 'Sections'
    hdr_row.cells[2].text = 'Enrolment'
    hdr_row.cells[3].merge(hdr_row.cells[4])
    hdr_row.cells[3].text = 'Score (Marka)'

    hdr_row2.cells[3].text = 'Markang ≥ 14'
    hdr_row2.cells[4].text = 'Markang < 14'

    SrpSetup.set_grade(table)

    SrpSetup.set_sections(table, all_sections)

    SrpSetup.set_total_students_per_grade(table, all_students)

    SrpSetup.set_total_per_section(table, all_students, all_sections)

    SrpSetup.set_school_total(table, all_students)

    TableSetupUtils.set_cell_center(table)

    TableSetupUtils.set_header_bold(table)
        
        
    # Save the document in-memory
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={school.name}School_Reading_Profile.docx'
    doc.save(response)

    return response