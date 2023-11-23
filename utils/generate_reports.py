from django.http import HttpResponse

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

from students.models import Student
from assessment.models import AssessmentSession
from read_app.models import School, ClassSection

from utils.isr_setup import IsrSetup
from utils.srp_setup import SrpSetup
from utils.table_setup_utils import TableSetupUtils


class GenerateReports():
    @staticmethod
    def generate_isr_word_document(student_id, language):
        student = Student.objects.get(pk=student_id)
        age = str(student.age)
        assessments = AssessmentSession.objects.filter(
            student=student, 
            assessment_type="Graded", 
            assessment_passage__passage__language=language,
            is_finished=True
        ).order_by('end_time')

        # Stop generating if there are no assessments finished
        if not assessments:
            return False

        doc = Document()

        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(11)

        # Add a centered heading with custom font and font size
        title = doc.add_paragraph('Individual Summary Record (ISR)')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_run = title.runs[0]
        title_run.font.size = Pt(13)
        title_run.bold = True
        
        title2 = doc.add_paragraph('Talaan ng Indibidwal na Pagbabasa (TIP)')
        title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title2_run = title2.runs[0]
        title2_run.font.size = Pt(13)
        title2_run.italic = True

        whitespace = doc.add_paragraph(' ')
        whitespace.runs[0].font.size = Pt(8.5)

        # Student Info
        info_line_1 = doc.add_paragraph('Name: ')
        info_line_1.add_run(f'{student.first_name} {student.last_name}').underline = True
        info_line_1.add_run('       Age: ')
        info_line_1.add_run(age).underline = True
        info_line_1.add_run(f'                          Grade/Section: ')
        info_line_1.add_run(f'{student.grade_level}-{student.class_section}').underline = True

        info_line_2 = doc.add_paragraph('School: ')
        info_line_2.add_run('Sapalibutad Elementary School').underline = True
        info_line_2.add_run('                           Teacher: ')
        info_line_2.add_run(f'{student.teacher}').underline = True

        if language == 'Filipino':
            info_line_3 = doc.add_paragraph('English: ☐ ')
            info_line_3.add_run('       Filipino: ☑')
        elif language == 'English':
            info_line_3 = doc.add_paragraph('English: ☑ ')
            info_line_3.add_run('       Filipino: ☐')
        
        for run in info_line_3.runs:
            run.bold = True


        # Table
        table = doc.add_table(rows=9, cols=10)
        table.style = 'Table Grid'
        hdr_row = table.rows[0]
        hdr2_row = table.rows[1]

        # Table Header First Layer 
        hdr_row.cells[3].merge(hdr_row.cells[5])
        hdr_row.cells[6].merge(hdr_row.cells[8])

        hdr_row.cells[0].text = 'Level Started'
        hdr_row.cells[1].text = 'Level'
        hdr_row.cells[2].text = 'Set'
        hdr_row.cells[3].text = 'Word Reading'
        hdr_row.cells[6].text = 'Comprehension'
        hdr_row.cells[9].text = 'Date Taken'

        # End First Layer

        # Table Header Second Layer
        hdr2_row.cells[0].text = 'Mark with an *'
        hdr2_row.cells[1].text = ''
        hdr2_row.cells[2].text = 'Indicate if A. B. C. or D.'
        hdr2_row.cells[3].text = 'Ind'
        hdr2_row.cells[4].text = 'Ins'
        hdr2_row.cells[5].text = 'Frus'
        hdr2_row.cells[6].text = 'Ind'
        hdr2_row.cells[7].text = 'Ins'
        hdr2_row.cells[8].text = 'Frus'
        hdr2_row.cells[9].text = ''
        # End Second Layer

        # Set the values for LEVEL column
        IsrSetup.set_level_column(table)

        IsrSetup.set_reading_rating(assessments, table)

        IsrSetup.set_level_started(assessments, table)

        TableSetupUtils.set_cell_center(table)

        TableSetupUtils.set_header_bold(table)

        legend = doc.add_paragraph('Legend: Ind- ')
        legend.runs[0].bold = True
        legend.add_run('Independent; ')
        legend.add_run('Ins- ').bold = True
        legend.add_run('Instructional; ')
        legend.add_run('Frus- ').bold = True
        legend.add_run('Frustration')

        # Add Line Break 
        doc.add_paragraph()

        page2_subtitle = doc.add_paragraph('Summary of Comprehension Responses')
        page2_subtitle.add_run('(Talaan ng Pag-unawa)')
        page2_subtitle.runs[0].bold = True


        # Table 
        page2_table = doc.add_table(rows=9, cols=12)
        page2_table.style = 'Table Grid'

        header = page2_table.rows[0]

        header.cells[0].text = 'Passage Level'
        header.cells[1].merge(header.cells[8])
        header.cells[1].text = 'Responses to Questions'
        header.cells[9].text = 'Score'
        header.cells[10].text = '%'
        header.cells[11].text = 'Reading Level'

        row3 = page2_table.rows[1]
        row3.cells[0].text = ' '

        for i in range(1, 9):
            row3.cells[i].text = f'Q{i}'

        IsrSetup.set_level_column(page2_table, col_index=0) 

        IsrSetup.set_passage_data(assessments, page2_table)

        TableSetupUtils.set_cell_center(page2_table)

        TableSetupUtils.set_header_bold(page2_table)

        # Save the document in-memory
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=Individual_Summary_Record_for_{student.last_name}_{student.first_name}.docx'
        doc.save(response)

        return response
    

    def generate_school_reading_profile(school_id):
        # Classes per grade and section
        school = School.objects.get(pk=school_id)
        all_students = Student.objects.all()
        all_sections = ClassSection.objects.all().order_by('section_name')

        # Setup
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(12)

        title = doc.add_paragraph('School Reading Profile')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(13)

        info_line1 = doc.add_paragraph('School: ')
        info_line1.add_run(school.name).underline = True
        info_line1.add_run("        Division: ")
        info_line1.add_run(school.division).underline = True
        info1_format = info_line1.paragraph_format
        info1_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        info_line2 = doc.add_paragraph('District: ')
        info_line2.add_run(school.district).underline = True
        info_line2.add_run("        Region: ")
        info_line2.add_run(school.region).underline = True
        info2_format = info_line2.paragraph_format
        info2_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        doc.add_paragraph()

        table = doc.add_table(rows=24, cols=5)
        table.style = 'Table Grid'
        hdr_row = table.rows[0]
        hdr_row2 = table.rows[1]

        hdr_row.cells[0].text = 'Grade'
        hdr_row.cells[1].text = 'Sections'
        hdr_row.cells[2].text = 'Enrolment'
        hdr_row.cells[3].merge(hdr_row.cells[4])
        hdr_row.cells[3].text = 'Score (Marka)'

        hdr_row2.cells[3].text = 'Markang ≥ 14'
        hdr_row2.cells[4].text = 'Markang < 14'

        SrpSetup.set_grade(table)

        SrpSetup.set_sections(table, all_sections)

        SrpSetup.set_total_students_per_grade(table, all_students)

        SrpSetup.set_total_per_section(table, all_students, all_sections)

        SrpSetup.set_school_total(table, all_students)

        TableSetupUtils.set_cell_center(table)

        TableSetupUtils.set_header_bold(table)
        
        
        # Save the document in-memory
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={school.name}School_Reading_Profile.docx'
        doc.save(response)

        return response